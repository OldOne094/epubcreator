"""اختبارات مستورد DOCX (M2.1): العناوين كفصول، تجاهل الجداول/الصور."""
from __future__ import annotations

import pytest

from app.core.importers import docx_to_chapters, docx_to_book, import_file


@pytest.fixture
def make_docx(tmp_path):
    def _make(name="book.docx", rows=None):
        from docx import Document

        doc = Document()
        for kind, text in rows or []:
            if kind == "h1":
                doc.add_heading(text, level=1)
            elif kind == "h2":
                doc.add_heading(text, level=2)
            elif kind == "p":
                doc.add_paragraph(text)
            elif kind == "table":
                t = doc.add_table(rows=2, cols=2)
                t.cell(0, 0).text = "خلية1"
                t.cell(1, 0).text = "خلية2"
        path = tmp_path / name
        doc.save(path)
        return path

    return _make


def test_docx_headings_become_chapters(make_docx):
    p = make_docx(
        rows=[
            ("h1", "الفصل الأول"),
            ("p", "نص أ."),
            ("h2", "الفصل الثاني"),
            ("p", "نص ب."),
        ]
    )
    chapters = docx_to_chapters(p)
    assert [c.title for c in chapters] == ["الفصل الأول", "الفصل الثاني"]
    assert chapters[0].body == "نص أ."
    assert chapters[1].body == "نص ب."


def test_docx_no_heading_single_chapter(make_docx):
    p = make_docx(rows=[("p", "سطر أول."), ("p", "سطر ثانٍ.")])
    chapters = docx_to_chapters(p)
    assert len(chapters) == 1
    assert "سطر أول." in chapters[0].body


def test_docx_tables_ignored(make_docx):
    p = make_docx(rows=[("h1", "فصل"), ("p", "نص"), ("table", None)])
    chapters = docx_to_chapters(p)
    assert "خلية" not in chapters[0].body


def test_docx_import_via_import_file(make_docx):
    p = make_docx(rows=[("h1", "بداية"), ("p", "محتوى.")])
    book = import_file(p)
    assert book.metadata.title == "book"
    assert [c.title for c in book.chapters] == ["بداية"]
    assert book.options.direction == "rtl"