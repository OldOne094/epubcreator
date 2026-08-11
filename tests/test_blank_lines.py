"""اختبارات بقاء الأسطر الفارغة بين الفقرات أثناء الاستيراد (مشكلة "الكومة الواحدة").

السطر الفارغ فاصل فقرات صريح: لا يجب أن يضيع في أي مسار استيراد، وإلا
ذاب النص في فقرة واحدة عند التحرير/التصدير.
"""
from __future__ import annotations

from pathlib import Path

import pytest

from app.core import importers
from app.core.clean import detect_chapters
from app.core.format import split_paragraphs


def test_detect_chapters_preserves_blank_lines():
    text = "الفصل الأول\n\nالفقرة الأولى.\n\nالفقرة الثانية."
    chunks = detect_chapters(text)
    assert chunks == ["الفصل الأول\n\nالفقرة الأولى.\n\nالفقرة الثانية."]


def test_detect_chapters_no_duplicate_blanks():
    text = "الفصل ١\n\n\nالفقرة الأولى.\n\n\n\nالفقرة الثانية."
    chunks = detect_chapters(text)
    assert "\n\n\n" not in chunks[0]


def test_txt_import_keeps_paragraph_separators(tmp_path):
    p = tmp_path / "كتاب.txt"
    p.write_text(
        "الفصل الأول\n\nالفقرة الأولى.\n\nالفقرة الثانية.",
        encoding="utf-8",
    )
    book = importers.import_file(p)
    body = book.chapters[0].body
    assert "الفقرة الأولى.\n\nالفقرة الثانية." in body
    # كل فقرة تُبنى فقرة مستقلة (وليس كومة واحدة)
    assert split_paragraphs(body) == ["الفقرة الأولى.", "الفقرة الثانية."]


def test_html_import_paragraphs_separated(tmp_path):
    p = tmp_path / "ك.html"
    p.write_text(
        '<html lang="ar"><body><h1>الفصل</h1>'
        "<p>فقرة أولى.</p><p>فقرة ثانية.</p><p>فقرة ثالثة.</p></body></html>",
        encoding="utf-8",
    )
    book = importers.import_file(p)
    assert book.chapters[0].body == "فقرة أولى.\n\nفقرة ثانية.\n\nفقرة ثالثة."


def test_html_import_long_paragraphs_stay_separate(tmp_path):
    long_a = "نص طويل " * 20
    long_b = "نص أطول " * 20
    p = tmp_path / "m.html"
    p.write_text(
        f'<body><h1>فصل</h1><p>{long_a}</p><p>{long_b}</p></body>',
        encoding="utf-8",
    )
    book = importers.import_file(p)
    paras = split_paragraphs(book.chapters[0].body)
    assert len(paras) == 2  # لا يدمجهما في فقرة واحدة


def test_markdown_import_keeps_separators(tmp_path):
    p = tmp_path / "k.md"
    p.write_text("# فصل\n\nأولى.\n\nثانية.", encoding="utf-8")
    book = importers.import_file(p)
    assert book.chapters[0].body == "أولى.\n\nثانية."


def test_docx_import_keeps_separators(tmp_path):
    from docx import Document

    p = tmp_path / "d.docx"
    doc = Document()
    doc.add_heading("الفصل", level=1)
    doc.add_paragraph("فقرة أولى.")
    doc.add_paragraph("فقرة ثانية.")
    doc.save(p)
    book = importers.import_file(p)
    assert book.chapters[0].body == "فقرة أولى.\n\nفقرة ثانية."


def test_rtf_import_keeps_separators(tmp_path):
    p = tmp_path / "r.rtf"
    p.write_bytes(
        "{\\rtf1\\ansi\\deff0 فقرة أولى.\\par\\par فقرة ثانية.}".encode("utf-8")
    )
    book = importers.import_file(p)
    assert "فقرة أولى." in book.chapters[0].body
    assert "فقرة ثانية." in book.chapters[0].body


@pytest.mark.parametrize(
    "body",
    [
        "فقرة أ.\n\nفقرة ب.\n\nفقرة ج.",
        "سطر قصير\n\nسطر قصير آخر\n\nثالث",
    ],
)
def test_roundtrip_editor_preserves_blanks(body):
    """الدورة عبر محرر QTextEdit لا تحذف الأسطر الفارغة."""
    from PySide6.QtWidgets import QApplication

    app = QApplication.instance() or QApplication([])
    from app.models import Book, Chapter
    from app.state import BookState
    from app.ui.pages import ChapterEditor

    b = Book()
    b.add_chapter(Chapter(title="فصل", body=body))
    editor = ChapterEditor(BookState(book=b))
    assert editor.body_edit.toPlainText() == body
    assert b.chapters[0].body == body
